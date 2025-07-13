from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()
doc.add_heading('FitNurture Advanced Posture Detection App - User Guide', 0)

doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'FitNurture Advanced is a posture detection application that allows users to analyze and record postural abnormalities using images. '
    'The advanced version provides the ability to fine-tune detection thresholds for various postural conditions, making it suitable for both general users and professionals who require more control over the analysis.'
)

doc.add_heading('Key Features', level=1)
doc.add_paragraph('- Analyze posture from single or multiple (4) views.')
doc.add_paragraph('- Adjust detection thresholds for conditions like Kyphosis, Lordosis, Tech Neck, Scoliosis, Flat Feet, Gait Abnormalities, Knock Knees, and Bow Legs.')
doc.add_paragraph('- Save and manage analysis results locally or upload to the cloud.')
doc.add_paragraph('- Generate PDF reports with images and recommendations.')
doc.add_paragraph('- Get AI-powered exercise and lifestyle suggestions.')

doc.add_heading('Getting Started', level=1)
doc.add_paragraph('1. Launch the App: Open the FitNurture Advanced app in your browser.')
doc.add_paragraph('2. Enter Student Details: Fill in the child\'s name (mandatory), select age group, gender, and indicate if the subject is wearing non-body-fitting clothes.')
doc.add_paragraph('3. Select Analysis Mode:')
doc.add_paragraph('   - Single View Analysis: For a quick check using a side view image.')
doc.add_paragraph('   - Multi-View Analysis (4 Views): For a comprehensive analysis using front, back, and both side views.')
doc.add_paragraph('4. Choose Input Method:')
doc.add_paragraph('   - Upload Image: Upload existing images.')
doc.add_paragraph('   - Use Camera: Capture images directly from your device.')

doc.add_heading('Adjusting Detection Thresholds', level=1)
doc.add_paragraph('Click the "⚙️ Advanced: Adjust Detection Thresholds" expander to reveal all adjustable parameters. Each threshold controls the sensitivity for detecting a specific postural abnormality.')

doc.add_heading('Example Thresholds and Their Effects', level=2)
thresholds = [
    ('Kyphosis (Sh-Hip Z-Diff >)', 'Higher value means only pronounced forward shoulder slouch is flagged.'),
    ('Lordosis (Hip-Knee Z-Diff >)', 'Higher value means only more pronounced lower back curve is flagged.'),
    ('Tech Neck (ESH Angle <)', 'Lower value means only severe forward head tilt is flagged.'),
    ('Tech Neck (ESH Horiz Dist >)', 'Higher value means ear must be much more forward of the shoulder to be flagged.'),
    ('Scoliosis (Shoulder Y-Diff >)', 'Higher value means only greater shoulder height differences are flagged.'),
    ('Flat Feet (Foot Arch <)', 'Lower value means only very flat arches are flagged.'),
    ('Gait Abnormality (Ankle X-Diff >)', 'Higher value means only wider foot stances are flagged.'),
    ('Knock Knees (Knee/Ankle Ratio <)', 'Lower value means knees must be much closer than ankles.'),
    ('Bow Legs (Knee/Ankle Ratio >)', 'Higher value means knees must be much wider than ankles.'),
]
for t, desc in thresholds:
    doc.add_paragraph(f'- {t}: {desc}')

doc.add_heading('How to Adjust', level=2)
doc.add_paragraph('Use the number input fields to set your desired threshold for each condition.')
doc.add_paragraph('Click "Reset Thresholds to Default" to restore original values.')
doc.add_paragraph('Adjusting thresholds allows you to make the detection stricter or more lenient based on your needs or the population being analyzed.')

doc.add_heading('Example Usage Scenarios', level=1)
doc.add_paragraph('Scenario 1: Screening in a School')
doc.add_paragraph('  - Use default thresholds for a general population.')
doc.add_paragraph('  - Capture images for all four views for each student.')
doc.add_paragraph('  - Analyze and save results locally or upload to the cloud for record-keeping.')
doc.add_paragraph('Scenario 2: Clinical Follow-Up')
doc.add_paragraph('  - Adjust thresholds to be stricter for follow-up on previously identified cases.')
doc.add_paragraph('  - Use single view analysis for quick checks.')
doc.add_paragraph('  - Generate PDF reports for patient records.')
doc.add_paragraph('Scenario 3: Research Study')
doc.add_paragraph('  - Fine-tune thresholds to match study criteria.')
doc.add_paragraph('  - Use multi-view analysis for comprehensive data.')
doc.add_paragraph('  - Export all results as CSV for further analysis.')

doc.add_heading('Saving and Managing Results', level=1)
doc.add_paragraph('Click "💾 Save Result Locally" to store the current analysis in your browser session.')
doc.add_paragraph('View all saved records in the data table below the analysis section.')
doc.add_paragraph('Download all records as CSV for offline use.')
doc.add_paragraph('Click "⬆️ Upload All Saved Records to Cloud" to store data in the Azure SQL database (requires configuration).')

doc.add_heading('Generating Reports', level=1)
doc.add_paragraph('Click "📄 Generate PDF Report" to create a detailed report including images, detected conditions, and recommendations.')
doc.add_paragraph('Download the PDF for sharing or record-keeping.')

doc.add_heading('AI-Powered Suggestions', level=1)
doc.add_paragraph('Click "✨ Get AI Exercise & Lifestyle Tips" to receive personalized advice based on detected conditions.')
doc.add_paragraph('Suggestions are generated using the Gemini AI model and are included in the PDF report if available.')

doc.add_heading('Tips for Best Results', level=1)
doc.add_paragraph('- Ensure good lighting and a clear, uncluttered background.')
doc.add_paragraph('- Subject should wear body-fitting clothes for more accurate analysis.')
doc.add_paragraph('- Capture the full body in each image.')
doc.add_paragraph('- Adjust thresholds as needed for your specific use case.')

doc.add_heading('Troubleshooting', level=1)
doc.add_paragraph('- If no person is detected, check image quality and ensure the full body is visible.')
doc.add_paragraph('- If results seem too strict or lenient, adjust the relevant thresholds.')
doc.add_paragraph('- For database upload issues, verify your cloud configuration and internet connection.')

doc.add_heading('Support', level=1)
doc.add_paragraph('For further assistance, refer to the in-app help or contact support at info@futurenurture.in.')

doc.save('FitNurture_Advanced_User_Guide.docx')
print('User guide generated as FitNurture_Advanced_User_Guide.docx')
